"""---------------------
- 이 파일이 담당하는 것:
  → 이력서 .txt 업로드, 텍스트 읽기, 맞춤 질문 생성 요청 준비,
    Function Calling 확인 데이터 표시, 질문 생성 대시보드.

- 이 파일이 담당하지 않는 것:
  → 면접 채팅 전체 화면 복사, API 키 입력, 8주차 agents 파일 수정.

- Day 5 self1로 넘길 값:
  → resume_file_name, resume_questions, resume_question_count, resume_step4b_done.
"""

import streamlit as st
from docx import Document
import io
import os
from openai import OpenAI


def read_resume_text(uploaded_file) -> str:
    """업로드된 이력서 파일에서 면접 질문 생성용 텍스트를 준비합니다."""
    #파일이 없으면 빈 문자열 반환
    if uploaded_file is None:
        return ""
    
    file_name = uploaded_file.name

    # .txt 파일 처리
    if file_name.endswith(".txt"):
        text = uploaded_file.read().decode("utf-8")

    # .docx 파일 처리
    elif file_name.endswith(".docx"):
        doc = Document(io.BytesIO(uploaded_file.read()))

        lines = []

        # 일반 단락 읽기
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                lines.append(paragraph.text)

        # 표 안의 텍스트 읽기
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text)

        text = "\n".join(lines)
    else:
        st.warning("지원하지 않는 파일 형식입니다.")
        return ""

    if len(text.strip()) < 50:
        st.warning("이력서 내용이 너무 짧습니다. 최소 50자 이상 입력해주세요.")
        return ""

    return text

def generate_questions_from_resume(
    resume_text: str,
    question_count: int,
    role_preset: str,
) -> list[str]:
    """이력서 텍스트로 실제 면접 질문을 생성합니다."""
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

    prompt = f"""
다음 이력서를 읽고 {role_preset} 관점에서 면접 질문 {question_count}개를 생성해주세요.

이력서:
{resume_text[:3000]}

규칙:
- 질문만 번호 없이 한 줄씩 작성
- 이력서 내용을 반영한 구체적인 질문
- 한국어로 작성
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": f"당신은 {role_preset} 면접관입니다."},
            {"role": "user", "content": prompt}
        ]
    )

    raw = response.choices[0].message.content
    questions = [
        line.strip()
        for line in raw.strip().split("\n")
        if line.strip()
    ]
    return questions[:question_count]


def build_resume_question_request(
    resume_text:str,
    question_count: int,
) -> dict:
    """이력서 기반 면접 질문 생성 요청 값을 만듭니다."""
    settings = st.session_state.get("settings", {})

    return{
        "resume_text": resume_text,
        "question_count": question_count,
        "model": settings.get("model", "gpt-4o-mini"),
        "system_prompt":settings.get(
            "system_prompt",
            "당신은 전문 면접관입니다."
        ),
        "role_preset": settings.get("role_preset", "기술 면접"),
    }

def render_function_call_result(result:dict) ->None:
    """질문 생성 결과와 도구 호출 확인 데이터를 분리해 표시합니다."""
    questions= result.get("questions", [])
    tool_calls = result.get("tool_calls", [])

    # 질문 목록 표시
    st.subheader("생성된 면접 질문")
    if questions:
        for i, q in enumerate(questions, 1):
            st.write(f"**{i}.** {q}")
    else:
        st.warning("생성된 질문이 없습니다.")

    # Function Calling 상세(접어두기)
    with st.expander("Function Calling 확인 데이터"):
        if tool_calls:
            st.json(tool_calls)
        else:
            st.write("호출 없음")


def save_resume_question_state(
        file_name:str,
        questions: list[str],
) -> None:
    """이력서 기반 질문 생성 결과를 세션 상태에 저장합니다."""
    st.session_state.resume_file_name = file_name
    st.session_state.resume_questions = questions
    st.session_state.resume_question_count = len(questions)
    st.session_state.resume_step4b_done=len(questions) > 0


def render_resume_dashboard()-> None:
    """이력서 기반 질문 생성 결과를 대시보드로 표시합니다."""
    questions = st.session_state.get("resume_questions", [])
    question_count = st.session_state.get("resume_question_count", 0)
    done = st.session_state.get("resume_step4b_done", False)
    file_name= st.session_state.get("resume_file_name", "없음")

    # 지표 표시
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("생성된 질문 수", question_count)
    with col2:
        st.metric("업로드 파일", file_name)

    with col3:
        st.metric("완료 상태", "완료" if done else "미완료")

    if questions:
        import pandas as pd

        def classify_length(q):
            if len(q) <= 30:
                return "단문 (30자 이하)"
            elif len(q) <= 60:
                return "중문 (32~60자)"
            else:
                return "장문 (60자 초과)"
        length_data = [classify_length(q) for q in questions]
        df = pd.DataFrame({"질문 길이": length_data})
        counts = df["질문 길이"].value_counts().reset_index()
        counts.columns = ["길이 분류", "개수"]
        st.bar_chart(counts.set_index("길이 분류"))

    # 진행률 표시( 목표 10문제 기준)
    progress = min(question_count/10, 1.0)
    st.progress(progress, text=f"목표 달성률: {int(progress * 100)}% (10문제 기준)")

# ============================
# 화면 구성
# ============================

st.title("이력서 분석")
st.caption(f"현재 질물 역할: {st.session_state.get('settings', {}).get('role_preset', '기술 면접')}")

# 파일 업로드
uploaded_file = st.file_uploader(
    "이력서 파일을 업로드하세요.",
    type=["txt", "docx"],
)

resume_text = read_resume_text(uploaded_file)

if resume_text:
    st.text_area(
        "이력서 미리보기 (앞 300자)",
        value=resume_text[:300] + "...",
        height=150,
        disabled=True,
    )
else:
    st.info("📂 .txt 또는 .docx 형식의 이력서 파일을 업로드해주세요.")

# 질문 수 선택
question_count = st.number_input(
    "생성할 질문 수",
    min_value=3,
    max_value=10,
    value=5,
    step=1,
)

# 질문 생성 버튼
if st.button("이력서 기반 질문 생성", disabled=not resume_text):
    with st.spinner("AI가 이력서를 분석하고 질문을 생성하고 있습니다..."):
        settings = st.session_state.get("settings", {})
        questions = generate_questions_from_resume(
            resume_text=resume_text,
            question_count=int(question_count),
            role_preset=settings.get("role_preset", "기술 면접"),
        )

    result = {
        "questions": questions,
        "tool_calls": []
    }

    save_resume_question_state(uploaded_file.name, questions)
    render_function_call_result(result)
    st.success(f"✅ {len(questions)}개 질문을 생성했습니다!")

# 면접 연습 시작 버튼 — 질문이 있을 때만 표시
if st.session_state.get("resume_questions"):
    if st.button("🎤 면접 연습 시작"):
        st.session_state.use_resume_questions = True
        st.session_state.resume_question_index = 0
        st.switch_page("pages/interview.py")

st.divider()
render_resume_dashboard()


# ============================
# Day 5 self1 연결 메모
# ============================
# - frontend/utils.py에서 반복되는 상태/오류 표시 함수를 정리한다
# - report.py에서 resume_questions를 읽어 결과 리포트를 만든다
# - README.md에 실행 순서와 포트 정보를 정리한다
# ============================
